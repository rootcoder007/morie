"""
Multi-format export engine for epidemiological analysis results.

Supports LaTeX, HTML, Markdown, Word (DOCX), Excel, figure export,
bundled results archives, analysis reports, reporting checklists
(CONSORT, STROBE), citation generation, and reproducibility manifests.
"""

from __future__ import annotations

import hashlib
import json
import logging
import os
import tempfile
import zipfile
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from typing import Any

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

logger = logging.getLogger(__name__)

# ===========================================================================
# DataFrame to LaTeX
# ===========================================================================


def df_to_latex(
    df: pd.DataFrame,
    *,
    caption: str = "",
    label: str = "",
    booktabs: bool = True,
    longtable: bool = False,
    sideways: bool = False,
    column_format: str | None = None,
    escape: bool = True,
    bold_header: bool = True,
    footnotes: list[str] | None = None,
    float_format: str = "%.3f",
) -> str:
    """
    Export a DataFrame to a LaTeX table string.

    Parameters
    ----------
    df : DataFrame
        Data to export.
    caption : str
        Table caption.
    label : str
        LaTeX label for cross-referencing.
    booktabs : bool
        Use booktabs rules (``\\toprule``, ``\\midrule``, ``\\bottomrule``).
    longtable : bool
        Use the ``longtable`` environment for multi-page tables.
    sideways : bool
        Wrap in ``sidewaystable`` for landscape orientation.
    column_format : str or None
        Column alignment string (e.g. ``"lccc"``).  Auto-generated if None.
    escape : bool
        Escape special LaTeX characters in cell values.
    bold_header : bool
        Bold column headers.
    footnotes : list of str or None
        Footnotes to append below the table.
    float_format : str
        Format string for floats.

    Returns
    -------
    str
        Complete LaTeX table source.
    """
    n_cols = len(df.columns) + (1 if df.index.name or True else 0)
    if column_format is None:
        column_format = "l" + "c" * len(df.columns)

    def _esc(s: str) -> str:
        if not escape:
            return s
        for ch in ["&", "%", "$", "#", "_", "{", "}"]:
            s = s.replace(ch, f"\\{ch}")
        return s

    lines: list[str] = []

    # Environment wrapper
    if sideways and not longtable:
        lines.append("\\begin{sidewaystable}[htbp]")
        lines.append("\\centering")
    elif not longtable:
        lines.append("\\begin{table}[htbp]")
        lines.append("\\centering")

    if caption:
        lines.append(f"\\caption{{{_esc(caption)}}}")
    if label:
        lines.append(f"\\label{{{label}}}")

    env = "longtable" if longtable else "tabular"
    lines.append(f"\\begin{{{env}}}{{{column_format}}}")

    if booktabs:
        lines.append("\\toprule")
    else:
        lines.append("\\hline")

    # Header
    idx_name = str(df.index.name) if df.index.name else ""
    header_cells = [_esc(idx_name)]
    for col in df.columns:
        cell = _esc(str(col))
        if bold_header:
            cell = f"\\textbf{{{cell}}}"
        header_cells.append(cell)
    lines.append(" & ".join(header_cells) + " \\\\")

    if booktabs:
        lines.append("\\midrule")
    else:
        lines.append("\\hline")

    # Rows
    for idx_val, row in df.iterrows():
        cells = [_esc(str(idx_val))]
        for val in row:
            if isinstance(val, float) and np.isfinite(val):
                cells.append(float_format % val)
            else:
                cells.append(_esc(str(val)))
        lines.append(" & ".join(cells) + " \\\\")

    if booktabs:
        lines.append("\\bottomrule")
    else:
        lines.append("\\hline")

    lines.append(f"\\end{{{env}}}")

    # Footnotes
    if footnotes:
        lines.append("\\begin{tablenotes}\\small")
        for fn in footnotes:
            lines.append(f"\\item {_esc(fn)}")
        lines.append("\\end{tablenotes}")

    if sideways and not longtable:
        lines.append("\\end{sidewaystable}")
    elif not longtable:
        lines.append("\\end{table}")

    return "\n".join(lines)


# ===========================================================================
# DataFrame to HTML
# ===========================================================================


def df_to_html(
    df: pd.DataFrame,
    *,
    title: str = "",
    css_classes: list[str] | None = None,
    responsive: bool = True,
    striped: bool = True,
    hover: bool = True,
    border: int = 0,
    float_format: str = "%.3f",
) -> str:
    """
    Export a DataFrame to an HTML table with styling.

    Parameters
    ----------
    df : DataFrame
    title : str
        Table caption / title.
    css_classes : list of str or None
        CSS classes to apply to the ``<table>`` element.
    responsive : bool
        Wrap in a responsive ``<div>``.
    striped : bool
        Apply striped-row styling.
    hover : bool
        Apply hover-row styling.
    border : int
        HTML border attribute.
    float_format : str
        Float formatting.

    Returns
    -------
    str
        Complete HTML table source.
    """
    classes = css_classes or []
    if striped:
        classes.append("table-striped")
    if hover:
        classes.append("table-hover")
    classes.append("morie-table")

    html = df.to_html(
        classes=classes,
        border=border,
        escape=False,
        float_format=lambda x: float_format % x if np.isfinite(x) else "",
    )

    parts = []
    if responsive:
        parts.append('<div class="table-responsive">')
    if title:
        parts.append(f'<h4 class="morie-table-title">{title}</h4>')
    parts.append(html)
    if responsive:
        parts.append("</div>")

    # Minimal embedded style
    style = """
<style>
.morie-table { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
  font-size: 13px; border-collapse: collapse; width: 100%; }
.morie-table th { background: #f8f9fa; border-bottom: 2px solid #dee2e6; padding: 8px 12px; text-align: left; }
.morie-table td { padding: 6px 12px; border-bottom: 1px solid #eee; }
.table-striped tbody tr:nth-child(odd) { background: #f8f9fa; }
.table-hover tbody tr:hover { background: #e9ecef; }
.morie-table-title { margin-bottom: 8px; color: #333; }
</style>
"""
    return style + "\n".join(parts)


# ===========================================================================
# DataFrame to Markdown
# ===========================================================================


def df_to_markdown(
    df: pd.DataFrame,
    *,
    title: str = "",
    float_format: str = "%.3f",
    index: bool = True,
) -> str:
    """
    Export a DataFrame to GitHub-Flavored Markdown.

    Parameters
    ----------
    df : DataFrame
    title : str
    float_format : str
    index : bool

    Returns
    -------
    str
    """
    # Format floats
    formatted = df.copy()
    for col in formatted.select_dtypes(include="number").columns:
        formatted[col] = formatted[col].apply(lambda x: float_format % x if np.isfinite(x) else "")

    md = formatted.to_markdown(index=index)
    parts = []
    if title:
        parts.append(f"**{title}**\n")
    parts.append(md)
    return "\n".join(parts)


# ===========================================================================
# DataFrame to Word/DOCX (optional dependency)
# ===========================================================================


def df_to_docx(
    df: pd.DataFrame,
    path: str,
    *,
    title: str = "",
    style: str = "Table Grid",
) -> str:
    """
    Export a DataFrame to a Word DOCX table.

    Requires ``python-docx``.  This is an optional dependency; an
    ``ImportError`` is raised with a helpful message if not installed.

    Parameters
    ----------
    df : DataFrame
    path : str
        Output file path.
    title : str
        Document / table title.
    style : str
        Word table style name.

    Returns
    -------
    str
        Path to the saved file.

    Raises
    ------
    ImportError
        If ``python-docx`` is not installed.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError:
        raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")

    doc = Document()
    if title:
        doc.add_heading(title, level=2)

    table = doc.add_table(rows=1, cols=len(df.columns) + 1)
    table.style = style

    # Header
    hdr = table.rows[0].cells
    hdr[0].text = str(df.index.name or "")
    for i, col in enumerate(df.columns):
        hdr[i + 1].text = str(col)
        for paragraph in hdr[i + 1].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for idx_val, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(idx_val)
        for i, val in enumerate(row):
            if isinstance(val, float) and np.isfinite(val):
                cells[i + 1].text = f"{val:.3f}"
            else:
                cells[i + 1].text = str(val)

    doc.save(path)
    logger.info("Saved DOCX table: %s", path)
    return path


# ===========================================================================
# DataFrame to Excel with formatting
# ===========================================================================


def df_to_excel(
    df: pd.DataFrame,
    path: str,
    *,
    sheet_name: str = "Results",
    title: str = "",
    bold_header: bool = True,
    alternating_rows: bool = True,
    number_format: str = "0.000",
    freeze_panes: tuple[int, int] | None = (1, 0),
) -> str:
    """
    Export a DataFrame to a formatted Excel file.

    Parameters
    ----------
    df : DataFrame
    path : str
        Output .xlsx file path.
    sheet_name : str
    title : str
        Optional title written above the table.
    bold_header : bool
    alternating_rows : bool
    number_format : str
    freeze_panes : tuple or None
        Row/column to freeze (default: freeze header row).

    Returns
    -------
    str
        Path to the saved file.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    start_row = 1
    if title:
        ws.cell(row=1, column=1, value=title).font = Font(bold=True, size=14)
        start_row = 3

    # Header
    header_font = Font(bold=True, color="FFFFFF") if bold_header else Font()
    header_fill = PatternFill(start_color="2c7bb6", end_color="2c7bb6", fill_type="solid")
    alt_fill = PatternFill(start_color="f0f4f8", end_color="f0f4f8", fill_type="solid")

    # Index header
    ws.cell(row=start_row, column=1, value=str(df.index.name or "")).font = header_font
    ws.cell(row=start_row, column=1).fill = header_fill
    ws.cell(row=start_row, column=1).alignment = Alignment(horizontal="center")

    for j, col in enumerate(df.columns):
        cell = ws.cell(row=start_row, column=j + 2, value=str(col))
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    # Data
    for i, (idx_val, row) in enumerate(df.iterrows()):
        r = start_row + 1 + i
        ws.cell(row=r, column=1, value=str(idx_val))
        if alternating_rows and i % 2 == 0:
            ws.cell(row=r, column=1).fill = alt_fill

        for j, val in enumerate(row):
            cell = ws.cell(row=r, column=j + 2)
            if isinstance(val, (int, float)) and np.isfinite(val):
                cell.value = float(val)
                cell.number_format = number_format
            else:
                cell.value = str(val)
            if alternating_rows and i % 2 == 0:
                cell.fill = alt_fill

    # Column widths
    for j in range(len(df.columns) + 1):
        ws.column_dimensions[chr(65 + j)].width = 15

    # Freeze panes
    if freeze_panes:
        ws.freeze_panes = ws.cell(row=start_row + freeze_panes[0], column=freeze_panes[1] + 1)

    wb.save(path)
    logger.info("Saved Excel file: %s", path)
    return path


# ===========================================================================
# Figure export to multiple formats
# ===========================================================================


def export_figure(
    fig: plt.Figure,
    base_path: str,
    *,
    formats: list[str] | None = None,
    dpi: int = 300,
    transparent: bool = False,
    close: bool = True,
) -> list[str]:
    """
    Save a matplotlib figure to multiple formats simultaneously.

    Parameters
    ----------
    fig : Figure
    base_path : str
        Base filename without extension (e.g. ``"output/figure1"``).
    formats : list of str or None
        File extensions.  Default: ``["pdf", "png", "svg"]``.
    dpi : int
        Resolution for raster formats.
    transparent : bool
        Transparent background.
    close : bool
        Close the figure after saving.

    Returns
    -------
    list of str
        Paths to all saved files.
    """
    if formats is None:
        formats = ["pdf", "png", "svg"]

    saved = []
    for fmt in formats:
        out = f"{base_path}.{fmt}"
        fig.savefig(
            out,
            format=fmt,
            dpi=dpi,
            bbox_inches="tight",
            pad_inches=0.05,
            transparent=transparent,
        )
        saved.append(out)
        logger.info("Saved figure: %s", out)

    if close:
        plt.close(fig)
    return saved


# ===========================================================================
# Results bundle (zip archive)
# ===========================================================================


def export_results_bundle(
    tables: dict[str, pd.DataFrame],
    figures: dict[str, plt.Figure] | None = None,
    metadata: dict[str, Any] | None = None,
    *,
    output_path: str = "results_bundle.zip",
    table_formats: list[str] | None = None,
    figure_formats: list[str] | None = None,
    dpi: int = 300,
) -> str:
    """
    Export all tables, figures, and metadata into a single ZIP archive.

    Parameters
    ----------
    tables : dict
        ``{name: DataFrame}`` mapping.
    figures : dict or None
        ``{name: Figure}`` mapping.
    metadata : dict or None
        Metadata dictionary (serialised as JSON).
    output_path : str
        Output zip file path.
    table_formats : list of str or None
        Formats for tables (default: ``["csv", "latex"]``).
    figure_formats : list of str or None
        Formats for figures (default: ``["pdf", "png"]``).
    dpi : int

    Returns
    -------
    str
        Path to the ZIP archive.
    """
    if table_formats is None:
        table_formats = ["csv", "latex"]
    if figure_formats is None:
        figure_formats = ["pdf", "png"]

    with tempfile.TemporaryDirectory() as tmpdir:
        # Tables
        tables_dir = os.path.join(tmpdir, "tables")
        os.makedirs(tables_dir)
        for name, df in tables.items():
            safe_name = name.replace(" ", "_").replace("/", "_")
            for fmt in table_formats:
                if fmt == "csv":
                    df.to_csv(os.path.join(tables_dir, f"{safe_name}.csv"))
                elif fmt == "latex":
                    latex_str = df_to_latex(df, caption=name)
                    with open(os.path.join(tables_dir, f"{safe_name}.tex"), "w") as f:
                        f.write(latex_str)
                elif fmt == "html":
                    html_str = df_to_html(df, title=name)
                    with open(os.path.join(tables_dir, f"{safe_name}.html"), "w") as f:
                        f.write(html_str)
                elif fmt == "markdown":
                    md_str = df_to_markdown(df, title=name)
                    with open(os.path.join(tables_dir, f"{safe_name}.md"), "w") as f:
                        f.write(md_str)

        # Figures
        if figures:
            figs_dir = os.path.join(tmpdir, "figures")
            os.makedirs(figs_dir)
            for name, fig in figures.items():
                safe_name = name.replace(" ", "_").replace("/", "_")
                export_figure(fig, os.path.join(figs_dir, safe_name), formats=figure_formats, dpi=dpi, close=False)

        # Metadata
        if metadata:
            with open(os.path.join(tmpdir, "metadata.json"), "w") as f:
                json.dump(metadata, f, indent=2, default=str)

        # Manifest
        manifest = {
            "created": datetime.now(timezone.utc).isoformat(),
            "tables": list(tables.keys()),
            "figures": list((figures or {}).keys()),
            "table_formats": table_formats,
            "figure_formats": figure_formats,
        }
        with open(os.path.join(tmpdir, "manifest.json"), "w") as f:
            json.dump(manifest, f, indent=2)

        # Zip
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for root, dirs, files in os.walk(tmpdir):
                for file in files:
                    fp = os.path.join(root, file)
                    arcname = os.path.relpath(fp, tmpdir)
                    zf.write(fp, arcname)

    logger.info("Saved results bundle: %s", output_path)
    return output_path


# ===========================================================================
# Analysis report generator
# ===========================================================================


def generate_report(
    title: str,
    sections: list[dict[str, Any]],
    *,
    output_path: str = "report.md",
    author: str = "",
    date: str | None = None,
) -> str:
    """
    Generate a Markdown analysis report from sections.

    Parameters
    ----------
    title : str
        Report title.
    sections : list of dict
        Each dict has keys: ``"heading"`` (str), ``"text"`` (str),
        and optionally ``"table"`` (DataFrame) and ``"figure"``
        (path to image).
    output_path : str
        Output file path.
    author : str
    date : str or None
        Date string.  Defaults to today.

    Returns
    -------
    str
        Path to the saved report.
    """
    if date is None:
        date = datetime.now().strftime("%Y-%m-%d")

    lines = [
        f"# {title}",
        "",
        f"**Author:** {author}" if author else "",
        f"**Date:** {date}",
        "",
        "---",
        "",
    ]

    for sec in sections:
        heading = sec.get("heading", "Section")
        text = sec.get("text", "")
        table = sec.get("table")
        figure = sec.get("figure")

        lines.append(f"## {heading}")
        lines.append("")
        if text:
            lines.append(text)
            lines.append("")
        if table is not None:
            lines.append(table.to_markdown())
            lines.append("")
        if figure is not None:
            lines.append(f"![{heading}]({figure})")
            lines.append("")

    content = "\n".join(lines)
    with open(output_path, "w") as f:
        f.write(content)

    logger.info("Saved report: %s", output_path)
    return output_path


# ===========================================================================
# CONSORT flow diagram data structure
# ===========================================================================


@dataclass
class ConsortNode:
    """A node in a CONSORT flow diagram."""

    label: str
    n: int
    excluded: int = 0
    exclusion_reasons: list[tuple[str, int]] | None = None


@dataclass
class ConsortDiagram:
    """CONSORT-style flow diagram data."""

    assessed: ConsortNode
    randomised: ConsortNode
    arms: list[dict[str, ConsortNode]]

    def to_dict(self) -> dict[str, Any]:
        """Serialise to a dictionary."""
        return {
            "assessed": asdict(self.assessed),
            "randomised": asdict(self.randomised),
            "arms": [{k: asdict(v) for k, v in arm.items()} for arm in self.arms],
        }

    def to_text(self) -> str:
        """Render as a text-based flow diagram."""
        lines = [
            f"Assessed for eligibility (n={self.assessed.n})",
        ]
        if self.assessed.excluded > 0:
            lines.append(f"  Excluded (n={self.assessed.excluded})")
            if self.assessed.exclusion_reasons:
                for reason, count in self.assessed.exclusion_reasons:
                    lines.append(f"    - {reason} (n={count})")
        lines.append(f"Randomised (n={self.randomised.n})")
        for arm in self.arms:
            for stage, node in arm.items():
                lines.append(f"  {stage}: {node.label} (n={node.n})")
                if node.excluded > 0:
                    lines.append(f"    Lost/excluded (n={node.excluded})")
                    if node.exclusion_reasons:
                        for reason, count in node.exclusion_reasons:
                            lines.append(f"      - {reason} (n={count})")
        return "\n".join(lines)


def create_consort_diagram(
    assessed_n: int,
    excluded_n: int,
    randomised_n: int,
    arms: list[dict[str, Any]],
    *,
    exclusion_reasons: list[tuple[str, int]] | None = None,
) -> ConsortDiagram:
    """
    Create a CONSORT flow diagram data structure.

    Parameters
    ----------
    assessed_n : int
        Number assessed for eligibility.
    excluded_n : int
        Number excluded.
    randomised_n : int
        Number randomised.
    arms : list of dict
        Each dict: ``{"label": str, "allocated": int,
        "analysed": int, "lost": int, ...}``.
    exclusion_reasons : list of (reason, count) or None

    Returns
    -------
    ConsortDiagram
    """
    assessed = ConsortNode(label="Assessed", n=assessed_n, excluded=excluded_n, exclusion_reasons=exclusion_reasons)
    randomised = ConsortNode(label="Randomised", n=randomised_n)

    diagram_arms = []
    for arm_data in arms:
        arm_dict = {}
        label = arm_data.get("label", "Arm")
        arm_dict["allocated"] = ConsortNode(
            label=f"Allocated to {label}",
            n=arm_data.get("allocated", 0),
        )
        arm_dict["analysed"] = ConsortNode(
            label=f"Analysed ({label})",
            n=arm_data.get("analysed", 0),
            excluded=arm_data.get("lost", 0),
        )
        diagram_arms.append(arm_dict)

    return ConsortDiagram(assessed=assessed, randomised=randomised, arms=diagram_arms)


# ===========================================================================
# STROBE checklist
# ===========================================================================


def strobe_checklist(
    study_type: str = "cross-sectional",
    items_addressed: dict[int, bool] | None = None,
) -> pd.DataFrame:
    """
    Generate a STROBE checklist for observational studies.

    Parameters
    ----------
    study_type : str
        ``"cross-sectional"``, ``"cohort"``, or ``"case-control"``.
    items_addressed : dict or None
        ``{item_number: True/False}`` indicating which items have been
        addressed.

    Returns
    -------
    DataFrame
        Checklist with item numbers, descriptions, and status.

    References
    ----------
    von Elm, E., et al. (2007). The Strengthening the Reporting of
    Observational Studies in Epidemiology (STROBE) Statement. *The
    Lancet*, 370(9596), 1453-1457.
    https://doi.org/10.1016/S0140-6736(07)61602-X
    """
    items = [
        (1, "Title and abstract", "Indicate study design; provide summary"),
        (2, "Background/rationale", "Scientific background and rationale"),
        (3, "Objectives", "State specific objectives and hypotheses"),
        (4, "Study design", "Present key elements of study design"),
        (5, "Setting", "Describe setting, locations, and dates"),
        (6, "Participants", "Eligibility criteria, sources, selection"),
        (7, "Variables", "Define outcomes, exposures, covariates"),
        (8, "Data sources", "Data sources and measurement methods"),
        (9, "Bias", "Describe efforts to address potential bias"),
        (10, "Study size", "Explain how study size was arrived at"),
        (11, "Quantitative variables", "Explain how variables were handled"),
        (12, "Statistical methods", "Describe all statistical methods"),
        (13, "Participants (results)", "Report numbers at each stage"),
        (14, "Descriptive data", "Characteristics of participants"),
        (15, "Outcome data", "Report numbers of outcome events"),
        (16, "Main results", "Unadjusted and adjusted estimates"),
        (17, "Other analyses", "Other analyses done (e.g. subgroup)"),
        (18, "Key results", "Summarise key results with objectives"),
        (19, "Limitations", "Discuss limitations and potential bias"),
        (20, "Interpretation", "Cautious interpretation of results"),
        (21, "Generalisability", "Discuss generalisability"),
        (22, "Funding", "Give source of funding and role of funders"),
    ]

    if items_addressed is None:
        items_addressed = {}

    records = []
    for num, title, description in items:
        records.append(
            {
                "Item": num,
                "Section": title,
                "Description": description,
                "Addressed": "Yes" if items_addressed.get(num, False) else "No",
            }
        )

    return pd.DataFrame(records)


# ===========================================================================
# Citation generator
# ===========================================================================


def generate_citation(
    title: str,
    authors: list[str],
    year: int,
    *,
    software_version: str = "0.1.0",
    doi: str = "",
    url: str = "https://github.com/hadesllm/morie",
    output_format: str = "bibtex",
) -> str:
    """
    Generate a citation entry for the analysis or the MORIE package.

    Parameters
    ----------
    title : str
    authors : list of str
        Author names.
    year : int
    software_version : str
    doi : str
    url : str
    output_format : str
        ``"bibtex"``, ``"apa"``, or ``"ris"``.

    Returns
    -------
    str
    """
    first_author_key = authors[0].split(",")[0].split()[-1].lower() if authors else "unknown"
    cite_key = f"{first_author_key}_{year}_morie"

    if output_format == "bibtex":
        author_str = " and ".join(authors)
        lines = [
            f"@software{{{cite_key},",
            f"  title = {{{title}}},",
            f"  author = {{{author_str}}},",
            f"  year = {{{year}}},",
            f"  version = {{{software_version}}},",
        ]
        if doi:
            lines.append(f"  doi = {{{doi}}},")
        lines.append(f"  url = {{{url}}},")
        lines.append("}")
        return "\n".join(lines)

    elif output_format == "apa":
        author_str = ", ".join(authors)
        s = f"{author_str} ({year}). {title}. Version {software_version}."
        if url:
            s += f" {url}"
        if doi:
            s += f" https://doi.org/{doi}"
        return s

    elif output_format == "ris":
        lines = [
            "TY  - COMP",
            f"TI  - {title}",
        ]
        for a in authors:
            lines.append(f"AU  - {a}")
        lines.append(f"PY  - {year}")
        if doi:
            lines.append(f"DO  - {doi}")
        if url:
            lines.append(f"UR  - {url}")
        lines.append("ER  -")
        return "\n".join(lines)

    raise ValueError(f"Unknown citation format: {output_format}")


# ===========================================================================
# Reproducibility manifest
# ===========================================================================


def create_reproducibility_manifest(
    data: pd.DataFrame | None = None,
    *,
    parameters: dict[str, Any] | None = None,
    seeds: dict[str, int] | None = None,
    output_path: str | None = None,
) -> dict[str, Any]:
    """
    Create a reproducibility manifest capturing environment and parameters.

    Parameters
    ----------
    data : DataFrame or None
        If provided, a SHA-256 checksum is computed.
    parameters : dict or None
        Analysis parameters.
    seeds : dict or None
        Random seeds.
    output_path : str or None
        If provided, save manifest as JSON.

    Returns
    -------
    dict
    """
    import sys

    manifest: dict[str, Any] = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "python_version": sys.version,
        "platform": sys.platform,
    }

    # Package versions
    versions: dict[str, str] = {}
    for pkg in ["pandas", "numpy", "scipy", "sklearn", "statsmodels", "matplotlib", "morie"]:
        try:
            mod = __import__(pkg)
            versions[pkg] = getattr(mod, "__version__", "unknown")
        except ImportError:
            pass
    manifest["package_versions"] = versions

    if data is not None:
        buf = data.to_csv(index=False).encode("utf-8")
        manifest["data_checksum_sha256"] = hashlib.sha256(buf).hexdigest()
        manifest["data_shape"] = list(data.shape)

    if parameters:
        manifest["parameters"] = parameters
    if seeds:
        manifest["random_seeds"] = seeds

    if output_path:
        with open(output_path, "w") as f:
            json.dump(manifest, f, indent=2, default=str)
        logger.info("Saved reproducibility manifest: %s", output_path)

    return manifest


# ===========================================================================
# JSON export for web visualisation
# ===========================================================================


def df_to_json(
    df: pd.DataFrame,
    path: str | None = None,
    *,
    orient: str = "records",
    date_format: str = "iso",
) -> str:
    """
    Export a DataFrame to JSON for web visualisation.

    Parameters
    ----------
    df : DataFrame
    path : str or None
        If provided, save to file.
    orient : str
        JSON orientation (default ``"records"``).
    date_format : str

    Returns
    -------
    str
        JSON string.
    """
    result = df.to_json(orient=orient, date_format=date_format, indent=2, default_handler=str)
    if path:
        with open(path, "w") as f:
            f.write(result)
        logger.info("Saved JSON: %s", path)
    return result


# ===========================================================================
# CSV with metadata sidecar
# ===========================================================================


def df_to_csv_with_meta(
    df: pd.DataFrame,
    path: str,
    *,
    description: str = "",
    source: str = "",
    variables: dict[str, str] | None = None,
) -> tuple[str, str]:
    """
    Export a DataFrame to CSV with a ``.csvmeta`` sidecar file.

    Parameters
    ----------
    df : DataFrame
    path : str
        Output CSV path.
    description : str
        Dataset description.
    source : str
        Data source.
    variables : dict or None
        ``{column_name: description}`` variable dictionary.

    Returns
    -------
    tuple of (csv_path, meta_path)
    """
    df.to_csv(path, index=True)

    meta_path = path + "meta"
    meta = {
        "file": os.path.basename(path),
        "description": description,
        "source": source,
        "created": datetime.now(timezone.utc).isoformat(),
        "shape": list(df.shape),
        "columns": list(df.columns),
        "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
    }
    if variables:
        meta["variable_descriptions"] = variables

    with open(meta_path, "w") as f:
        json.dump(meta, f, indent=2)

    logger.info("Saved CSV: %s  Meta: %s", path, meta_path)
    return path, meta_path


# ===========================================================================
# Clipboard export
# ===========================================================================


def to_clipboard(
    df: pd.DataFrame,
    *,
    fmt: str = "text",
    sep: str = "\t",
) -> None:
    """
    Copy a DataFrame to the system clipboard for pasting.

    Parameters
    ----------
    df : DataFrame
    fmt : str
        ``"text"`` (tab-separated) or ``"html"``.
    sep : str
        Separator for text format.
    """
    if fmt == "html":
        html = df.to_html()
        df.to_clipboard(excel=True)
        logger.info("Copied HTML table to clipboard")
    else:
        df.to_clipboard(sep=sep)
        logger.info("Copied table to clipboard (sep='%s')", sep)


# ===========================================================================
# LaTeX preamble generator
# ===========================================================================


def latex_preamble(
    *,
    booktabs: bool = True,
    longtable: bool = False,
    sideways: bool = False,
    amsmath: bool = True,
    hyperref: bool = True,
    natbib: bool = True,
    threeparttable: bool = True,
) -> str:
    """
    Generate a LaTeX preamble with required packages for MORIE tables.

    Parameters
    ----------
    booktabs : bool
    longtable : bool
    sideways : bool
    amsmath : bool
    hyperref : bool
    natbib : bool
    threeparttable : bool

    Returns
    -------
    str
    """
    packages = []
    if booktabs:
        packages.append("\\usepackage{booktabs}")
    if longtable:
        packages.append("\\usepackage{longtable}")
    if sideways:
        packages.append("\\usepackage{rotating}")
    if amsmath:
        packages.append("\\usepackage{amsmath}")
    if hyperref:
        packages.append("\\usepackage{hyperref}")
    if natbib:
        packages.append("\\usepackage{natbib}")
    if threeparttable:
        packages.append("\\usepackage{threeparttable}")

    packages.extend(
        [
            "\\usepackage{graphicx}",
            "\\usepackage[utf8]{inputenc}",
            "\\usepackage[T1]{fontenc}",
        ]
    )

    return "\n".join(packages)


# ===========================================================================
# Batch export
# ===========================================================================


def batch_export(
    results: dict[str, pd.DataFrame | plt.Figure],
    output_dir: str,
    *,
    table_formats: list[str] | None = None,
    figure_formats: list[str] | None = None,
    dpi: int = 300,
    manifest: bool = True,
) -> list[str]:
    """
    Batch export all tables and figures from a pipeline run.

    Parameters
    ----------
    results : dict
        ``{name: DataFrame_or_Figure}`` mapping.
    output_dir : str
        Output directory.
    table_formats : list of str or None
        Default: ``["csv", "latex"]``.
    figure_formats : list of str or None
        Default: ``["pdf", "png"]``.
    dpi : int
    manifest : bool
        Write a JSON manifest of exported files.

    Returns
    -------
    list of str
        Paths to all exported files.
    """
    if table_formats is None:
        table_formats = ["csv", "latex"]
    if figure_formats is None:
        figure_formats = ["pdf", "png"]

    os.makedirs(output_dir, exist_ok=True)
    tables_dir = os.path.join(output_dir, "tables")
    figs_dir = os.path.join(output_dir, "figures")
    os.makedirs(tables_dir, exist_ok=True)
    os.makedirs(figs_dir, exist_ok=True)

    exported: list[str] = []

    for name, obj in results.items():
        safe_name = name.replace(" ", "_").replace("/", "_")

        if isinstance(obj, pd.DataFrame):
            for fmt in table_formats:
                if fmt == "csv":
                    path = os.path.join(tables_dir, f"{safe_name}.csv")
                    obj.to_csv(path)
                    exported.append(path)
                elif fmt == "latex":
                    path = os.path.join(tables_dir, f"{safe_name}.tex")
                    with open(path, "w") as f:
                        f.write(df_to_latex(obj, caption=name))
                    exported.append(path)
                elif fmt == "html":
                    path = os.path.join(tables_dir, f"{safe_name}.html")
                    with open(path, "w") as f:
                        f.write(df_to_html(obj, title=name))
                    exported.append(path)
                elif fmt == "markdown":
                    path = os.path.join(tables_dir, f"{safe_name}.md")
                    with open(path, "w") as f:
                        f.write(df_to_markdown(obj, title=name))
                    exported.append(path)

        elif isinstance(obj, plt.Figure):
            paths = export_figure(
                obj,
                os.path.join(figs_dir, safe_name),
                formats=figure_formats,
                dpi=dpi,
                close=True,
            )
            exported.extend(paths)

    if manifest:
        manifest_data = {
            "exported_at": datetime.now(timezone.utc).isoformat(),
            "files": exported,
            "n_tables": sum(1 for v in results.values() if isinstance(v, pd.DataFrame)),
            "n_figures": sum(1 for v in results.values() if isinstance(v, plt.Figure)),
        }
        manifest_path = os.path.join(output_dir, "export_manifest.json")
        with open(manifest_path, "w") as f:
            json.dump(manifest_data, f, indent=2)
        exported.append(manifest_path)

    logger.info("Batch exported %d files to %s", len(exported), output_dir)
    return exported
